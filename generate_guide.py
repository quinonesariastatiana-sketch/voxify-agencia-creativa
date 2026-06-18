"""Genera la Guía de Configuración de Marca en Voxify como documento Word."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ───────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0A, 0x25, 0x40)
PURPLE = RGBColor(0x63, 0x5B, 0xFF)
GREEN  = RGBColor(0x05, 0x96, 0x69)
ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
RED    = RGBColor(0xEF, 0x44, 0x44)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF4, 0xF6, 0xFA)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in borders.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = color
    return p

def heading2(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color
    return p

def heading3(text, color=PURPLE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = color
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def check(text, indent=0.4):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(indent)
    # Checkbox symbol
    cb = p.add_run("☐  ")
    cb.font.size = Pt(10)
    cb.font.color.rgb = PURPLE
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def note(text, color=GRAY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.6)
    run = p.add_run("💡  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.italic = True
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)

def banner(text, bg="0A2540", fg=WHITE):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, bg)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = fg
    doc.add_paragraph()

def simple_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        shade_cell(cell, "0A2540")
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
    # Data rows
    for ri, row in enumerate(rows):
        bg = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            shade_cell(cell, bg)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                tbl.cell(ri, ci).width = Cm(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
cell = tbl.cell(0, 0)
shade_cell(cell, "0A2540")
cell.paragraphs[0].clear()
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(32)
p.paragraph_format.space_after  = Pt(32)
r1 = p.add_run("VOXIFY\n")
r1.bold = True; r1.font.size = Pt(36); r1.font.color.rgb = WHITE
r2 = p.add_run("Guía de Configuración de Marca\n")
r2.font.size = Pt(18); r2.font.color.rgb = RGBColor(0x9B, 0x7C, 0xFF)
r3 = p.add_run("Lista de chequeo paso a paso — sin comandos ni tecnicismos")
r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Para cualquier persona · Sin conocimientos técnicos · Junio 2026")
r.font.size = Pt(10); r.font.color.rgb = GRAY; r.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ══════════════════════════════════════════════════════════════════════════
heading1("Contenido de esta guía")
index_items = [
    ("Antes de empezar",              "Cuentas y preparación necesaria"),
    ("Parte 1 — Crear la marca",      "Completar las 8 pestañas del panel Voxify"),
    ("Parte 2 — Conectar Facebook e Instagram", "Tokens de Meta paso a paso"),
    ("Parte 3 — Conectar LinkedIn",   "App y token de LinkedIn (opcional)"),
    ("Parte 4 — Logo, Web y Redes",   "Subir logo, manual y analizar con IA"),
    ("Parte 5 — Configurar el agente","Sistema de IA y estrategia"),
    ("Parte 6 — Primer contenido",    "Aprobar y publicar los primeros posts"),
    ("Lista de chequeo final",        "Verificación completa antes de lanzar"),
    ("Solución de problemas",         "Errores comunes y cómo resolverlos"),
]
simple_table(["Sección", "Qué cubre"], index_items, col_widths=[7, 10])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  ANTES DE EMPEZAR
# ══════════════════════════════════════════════════════════════════════════
banner("ANTES DE EMPEZAR — Lo que necesitas tener listo")
heading2("Cuentas obligatorias")
check("Cuenta de Facebook personal — la que usarás para administrar todo")
check("Página de Facebook para la marca (no un perfil personal)")
check("Cuenta de Instagram Business o Creador — conectada a esa Página de Facebook")
check("Acceso de Administrador en ambas (Página y Business Manager)")
check("Clave de API de Anthropic — para que funcione el agente de IA")

heading2("Cómo conectar Instagram a tu Página de Facebook")
note("Hazlo antes de generar los tokens. Si ya están conectados, salta este paso.")
check("Abre Instagram en el celular → ve a tu perfil")
check("Toca el menú ☰ (tres líneas) → Configuración y actividad")
check("Toca Cuenta → Cambiar tipo de cuenta → selecciona Cuenta profesional")
check("Toca Cuenta → Cuentas vinculadas → selecciona Facebook")
check("Inicia sesión con tu Facebook y elige la Página de esa marca")
check("Confirma. Ya están conectadas.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  PARTE 1 — CREAR LA MARCA
# ══════════════════════════════════════════════════════════════════════════
banner("PARTE 1 — Crear la marca en Voxify")
check("Abre el navegador y ve a:  http://localhost:5000/brands")
check("Haz clic en el botón morado  + Nueva Marca  en el panel izquierdo")
check("Se abrirá un formulario con 8 pestañas. Completa cada una en orden.")

divider()
heading2("Pestaña 1 — 🏷 Identidad")
simple_table(
    ["Campo", "Qué poner"],
    [
        ("ID de Marca (slug)", "Una palabra sin espacios ni mayúsculas. Ej: aroma_beans"),
        ("Nombre", "El nombre completo. Ej: Aroma Beans"),
        ("Tagline / Eslogan", "La frase corta. Ej: Two coffees. Two stories. One Colombia."),
        ("Color principal", "Haz clic en el cuadro de color y elige el color de la marca"),
        ("Color del texto", "Blanco si el fondo es oscuro, Negro si es claro"),
        ("Industria / Sector", "Ej: Specialty Coffee / E-commerce"),
        ("Geografía objetivo", "Dónde opera la marca. Ej: Miami FL, Austin TX, Portland OR"),
        ("Descripción", "2-3 frases: qué hace, a quién sirve, qué la hace especial"),
        ("Misión", "La razón de ser más profunda de la marca (más allá del dinero)"),
        ("Valores", "Escribe cada valor y presiona Enter. Ej: Autenticidad, Comunidad"),
        ("Hashtags", "Escribe cada hashtag y presiona Enter. Ej: #ColombiaCoffee"),
    ],
    col_widths=[5.5, 11.5]
)

divider()
heading2("Pestaña 2 — 👥 Audiencia")
check("Idioma del contenido: selecciona Español o Inglés")
check("Canales activos: marca las casillas de Instagram, Facebook, LinkedIn")
check("Geografía de la audiencia: ciudad o país del público objetivo")
check("Personas: describe 1 a 3 tipos de cliente ideal, cada uno con:")
body("• Nombre descriptivo  •  Rango de edad  •  Ocupación", indent=1.2)
body("• Dolor principal (qué problema tiene que tu marca resuelve)", indent=1.2)
body("• Objetivo (qué quiere lograr con tu producto/servicio)", indent=1.2)
note("Piensa en un cliente real. Descríbelo como si le contaras a alguien cómo es esa persona.")

divider()
heading2("Pestaña 3 — 🎙 Voz & Tono")
check("Adjetivos de voz: palabras que describen el tono. Ej: cercano, educado, directo")
check("Qué NUNCA decir: palabras o frases prohibidas. Ej: premium, artesanal con amor")
check("Ejemplo de texto BUENO: pega un post que suene como quieres que hable la marca")
check("Ejemplo de texto MALO: pega un ejemplo de lo que NO quieres")
check("Nivel de formalidad: desliza entre informal (0) y formal (1)")
check("Uso de emojis: selecciona ninguno, moderado o frecuente")

divider()
heading2("Pestaña 4 — 🎯 Posicionamiento")
check("Propuesta única de valor (USP): por qué tu marca es diferente a todo lo demás")
check("Competidores: agrega 2-3 con su nombre y su punto débil principal")
check("Diferenciadores: lo que hace única a tu marca. Mínimo 3, máximo 7")

divider()
heading2("Pestaña 5 — 🔑 Credenciales")
body("Déjala vacía por ahora. Se completa en la Parte 2 de esta guía.", indent=0.5)

divider()
heading2("Pestaña 6 — 📅 Calendario de publicación")
check("Marca los días de la semana en que quieres publicar en cada plataforma")
check("Selecciona la hora de publicación (en Eastern Time — ET)")
note("Recomendaciones probadas:")
simple_table(
    ["Plataforma", "Días sugeridos", "Hora (ET)"],
    [
        ("📸 Instagram", "Martes a Viernes", "9:00 AM"),
        ("👍 Facebook",  "Lunes a Viernes",  "10:00 AM"),
        ("💼 LinkedIn",  "Solo Miércoles (si aplica)", "8:00 AM"),
    ],
    col_widths=[4.5, 7, 5.5]
)
check("Si no usas LinkedIn, deja todos sus días sin marcar")

divider()
heading2("Pestaña 7 — 📊 Objetivos KPI")
check("Completa las metas a 30, 60 y 90 días")
simple_table(
    ["Métrica", "30 días", "60 días", "90 días"],
    [
        ("Seguidores Instagram", "____", "____", "____"),
        ("Engagement rate %",    "3-5%", "4-5%", "5%+"),
        ("Alcance Facebook",     "____", "____", "____"),
        ("Leads / emails",       "____", "____", "____"),
        ("Clientes",             "____", "____", "____"),
        ("Ingresos USD",         "$___", "$___", "$___"),
    ],
    col_widths=[5.5, 4, 4, 4]
)
check("Haz clic en  Guardar Marca  (botón verde arriba a la derecha)")

divider()
heading2("Pestaña 8 — 🚀 Estrategia  (después de guardar)")
check("Una vez guardada la marca, regresa a esta pestaña")
check("Haz clic en  Generar Plan 90 Días  — espera 1-2 minutos")
check("Haz clic en  Actualizar Plan Mensual")
check("Haz clic en  Generar Semana  — crea los primeros 10-11 posts")
note("Al terminar cada generación, el panel cambia automáticamente a la pestaña Estrategia y muestra el resultado.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  PARTE 2 — CONECTAR META
# ══════════════════════════════════════════════════════════════════════════
banner("PARTE 2 — Conectar Facebook e Instagram (Meta)")
body("Esta es la parte más técnica. Sigue cada paso exactamente.")

heading2("Paso A — Entrar a Meta for Developers")
check("Ve a:  developers.facebook.com")
check("Inicia sesión con tu cuenta personal de Facebook (la que administra la Página)")
check("Haz clic en  My Apps  (arriba a la derecha) →  Create App")

heading2("Paso B — Crear la aplicación de Meta")
check("Te pregunta el tipo de app. Selecciona  Other  →  Next")
check("Tipo de app: selecciona  Business  →  Next")
check("Escribe un nombre cualquiera (Ej: VoxifyPublisher) →  Create App")
check("Es posible que te pida confirmar tu contraseña")
check("Ya estás dentro del panel de tu nueva app")

heading2("Paso C — Agregar los permisos necesarios")
check("En el menú izquierdo busca  App Review  →  Permissions and Features")
check("Busca y solicita (haz clic en Request o Add en cada uno):")
simple_table(
    ["Permiso", "Para qué sirve"],
    [
        ("instagram_basic",           "Leer info básica del perfil"),
        ("instagram_content_publish", "Publicar posts en Instagram"),
        ("pages_manage_posts",        "Crear y publicar en tu Página de Facebook"),
        ("pages_read_engagement",     "Leer estadísticas de la página"),
        ("pages_show_list",           "Ver la lista de páginas que administras"),
        ("business_management",       "Acceso a Business Manager"),
    ],
    col_widths=[6, 11]
)
check("Anota el  App ID  y el  App Secret  (Configuración → Básico)")

heading2("Paso D — Obtener el Page Access Token")
check("Ve a:  developers.facebook.com/tools/explorer")
check("Arriba a la derecha, en el menú  Meta App, selecciona tu app")
check("Haz clic en  Generate Access Token  — aparecerá una ventana de Facebook")
check("Dale permisos a todo lo que te pide → continúa")
check("Ya tienes un  User Access Token  (empieza con EAA...) — es temporal, 1 hora")

heading2("Paso E — Encontrar tu Page ID y Page Token")
check("Con ese token activo, en el campo URL del Explorer escribe:  /me/accounts")
check("Haz clic en el botón azul  Submit")
check("Verás una lista de tus páginas. Busca la página de la marca y copia:")
body("→  El número en el campo  id  = tu Facebook Page ID", indent=1.2)
body("→  El valor en  access_token  = tu Page Access Token", indent=1.2)
check("Guarda ambos valores en un documento de texto")

heading2("Paso F — Encontrar tu Instagram Business Account ID")
check("En el Explorer, cambia la URL a:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r = p.add_run("/{TU-PAGE-ID}?fields=instagram_business_account")
r.font.name = "Courier New"; r.font.size = Pt(9.5)
r.font.color.rgb = PURPLE
check("Reemplaza {TU-PAGE-ID} con el número que copiaste")
check("Haz clic en  Submit")
check('Verás: "instagram_business_account": {"id": "17841xxxxxxxxx"}')
check("Ese número es tu  Instagram Business Account ID")

heading2("Paso G — Hacer el token de larga duración (60 días)")
check("En el Explorer, escribe esta URL (reemplaza los 3 valores):")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r = p.add_run("/oauth/access_token?grant_type=fb_exchange_token\n"
              "&client_id={APP-ID}&client_secret={APP-SECRET}\n"
              "&fb_exchange_token={PAGE-TOKEN}")
r.font.name = "Courier New"; r.font.size = Pt(9)
r.font.color.rgb = PURPLE
check("Haz clic en Submit → copia el access_token del resultado")
check("Este nuevo token dura 60 días")
note("Alternativa sin vencimiento: En Meta Business Manager → Configuración → Usuarios del sistema → crea un System User → genera un token permanente.")

heading2("Paso H — Ingresar las credenciales en Voxify")
check("Regresa a  http://localhost:5000/brands")
check("Selecciona tu marca → pestaña  🔑 Credenciales")
check("Pega el  Instagram Business Account ID  en su campo")
check("Pega el  Facebook Page ID  en su campo")
check("Pega el  Page Access Token  en el campo  Facebook Page Access Token")
check("Pega el mismo token en  Meta System User Token")
check("Haz clic en  Probar IG  → debe aparecer un mensaje verde de éxito")
check("Haz clic en  Probar FB  → mismo resultado")
check("Haz clic en  Guardar Marca")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  PARTE 3 — LINKEDIN
# ══════════════════════════════════════════════════════════════════════════
banner("PARTE 3 — Conectar LinkedIn  (opcional — solo para marcas B2B)")

heading2("Paso A — Crear la app en LinkedIn")
check("Ve a:  linkedin.com/developers")
check("Inicia sesión → haz clic en  Create App")
check("App name: VoxifyPublisher")
check("LinkedIn Page: busca y selecciona la Página de tu empresa")
check("App logo: sube cualquier imagen → Create App")

heading2("Paso B — Solicitar permisos")
check("Ve a la pestaña  Products  dentro de tu app")
check("Solicita:  Share on LinkedIn  →  Request access")
check("Solicita:  Marketing Developer Platform  →  Request access")
note("LinkedIn revisa las solicitudes en 1-3 días hábiles.")

heading2("Paso C — Obtener el Access Token")
check("Una vez aprobado, ve a la pestaña  Auth  dentro de tu app")
check("En la sección  OAuth 2.0 Tools, haz clic en  OAuth Token Generator")
check("Selecciona los permisos: w_member_social, w_organization_social")
check("Haz clic en  Request access token  → acepta la autorización")
check("Copia el  Access Token  resultante")

heading2("Paso D — Obtener el Organization ID")
check("Ve a la página de tu empresa en LinkedIn")
check("Mira la URL:  linkedin.com/company/XXXXXXXX/")
check("Ese número  XXXXXXXX  es tu  LinkedIn Organization ID")

heading2("Paso E — Ingresar en Voxify")
check("Regresa a la pestaña  🔑 Credenciales  de tu marca")
check("Pega el  LinkedIn Access Token  y el  Organization ID")
check("Guarda la marca")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  PARTE 4 — LOGO, WEB Y REDES
# ══════════════════════════════════════════════════════════════════════════
banner("PARTE 4 — Logo, Sitio Web, Redes Sociales y Manual de Marca")

heading2("Subir el logo")
check("Ve a la pestaña  🏷 Identidad")
check("Busca la sección  Identidad Visual")
check("Arrastra el archivo del logo a la zona punteada, o haz clic en ella para buscarlo")
check("Formatos aceptados: PNG, JPG, SVG, WEBP — máx 5 MB")
check("Verás el preview inmediatamente. El logo se guarda automáticamente.")

heading2("Agregar presencia digital")
check("En la misma pestaña Identidad, baja hasta la sección  🌐 Presencia Digital")
check("Escribe la URL del sitio web:  https://tumarca.com")
check("Agrega los handles o URLs de redes sociales:")
body("• Instagram: @tumarca o la URL completa del perfil", indent=1.2)
body("• Facebook: @tumarca o URL de la página", indent=1.2)
body("• LinkedIn: linkedin.com/company/tumarca", indent=1.2)
body("• TikTok: @tumarca (si aplica)", indent=1.2)
body("• YouTube y Tienda: si aplica", indent=1.2)

heading2("Subir el manual de marca")
check("Baja hasta la sección  📋 Manual de Marca y Recursos")
check("Arrastra el PDF del brand guidelines, o haz clic para buscarlo")
check("Formatos aceptados: PDF, TXT, MD — máx 10 MB")
check("El servidor extrae el texto automáticamente y lo usa para el análisis")

heading2("Analizar la marca con IA")
check("Con la web y/o el manual cargados, haz clic en:  🤖 Analizar marca con IA")
check("Claude leerá el sitio web, las URLs de redes y el manual en 30-60 segundos")
check("Aparecerá un panel verde con todos los campos sugeridos:")
body("• Descripción, misión, industria, geografía", indent=1.2)
body("• Valores, hashtags, voz y tono", indent=1.2)
body("• Propuesta única de valor y diferenciadores", indent=1.2)
body("• Audiencia y personas de cliente ideal", indent=1.2)
body("• Pilares de contenido con porcentajes", indent=1.2)
check("Revisa el resumen y haz clic en  ✓ Aplicar todos los campos sugeridos")
check("El formulario se llena automáticamente. Ajusta lo que necesites y guarda.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  PARTE 5 — AGENTE
# ══════════════════════════════════════════════════════════════════════════
banner("PARTE 5 — Configurar el Agente de IA")

heading2("Verificar el System Prompt")
check("Ve a la pestaña  🚀 Estrategia")
check("El campo  Estrategia 90 Días  debe contener las instrucciones del agente")
check("Si está vacío, usa el botón  Generar Prompt  o escríbelo manualmente")
check("El prompt debe mencionar: nombre de marca, tono de voz, qué nunca decir, audiencia")

heading2("Generar la estrategia completa")
check("Haz clic en  Generar Plan 90 Días  — espera 1-2 minutos")
check("Lee el plan generado. Verifica coherencia con tus objetivos")
check("Si algo no cuadra, ajusta la descripción o misión (pestaña Identidad) y vuelve a generar")
check("Haz clic en  Actualizar Plan Mensual")
check("Haz clic en  ⚡ Generar Semana  — el agente crea los primeros 10-11 posts")
note("El contenido semanal aparecerá en la pantalla de revisión (http://localhost:5000)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  PARTE 6 — PRIMER CONTENIDO
# ══════════════════════════════════════════════════════════════════════════
banner("PARTE 6 — Aprobar y publicar el primer contenido")
check("Ve a  http://localhost:5000  (pantalla de revisión principal)")
check("En la esquina superior derecha, selecciona tu marca en el menú desplegable")
check("Verás los posts generados en estado  Pendiente de aprobación")
body("Por cada post:")
check("Léelo completo — verifica que suene a la marca", indent=0.8)
check("Si está bien: haz clic en  Aprobar  (queda listo para su horario)", indent=0.8)
check("Si necesita cambio: edita el texto directamente y luego aprueba", indent=0.8)
check("Si no sirve: haz clic en  Rechazar", indent=0.8)
check("Con al menos 1 post aprobado, el scheduler publica automáticamente en el horario configurado")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  LISTA DE CHEQUEO FINAL
# ══════════════════════════════════════════════════════════════════════════
banner("LISTA DE CHEQUEO FINAL — Antes de declarar la marca lista")
check("Información de identidad completa (nombre, tagline, descripción, misión, valores)")
check("Logo subido y visible en el panel")
check("Sitio web y redes sociales registradas")
check("Manual de marca cargado (si existe)")
check("Análisis con IA aplicado y campos revisados")
check("Al menos 2 personas de audiencia definidas")
check("Voz y tono con ejemplos de qué SÍ y qué NO")
check("USP y diferenciadores claros")
check("Credenciales probadas y con check verde (IG + FB)")
check("Calendario de publicación configurado")
check("Metas a 30/60/90 días definidas")
check("Plan de 90 días generado y visible")
check("Plan mensual generado")
check("Al menos 5 posts aprobados y listos para publicar")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  SOLUCIÓN DE PROBLEMAS
# ══════════════════════════════════════════════════════════════════════════
banner("SOLUCIÓN DE PROBLEMAS COMUNES")
simple_table(
    ["Síntoma", "Causa probable", "Solución"],
    [
        ("'Probar IG' sale error rojo",
         "El token venció o no tiene permisos",
         "Vuelve al Graph API Explorer y genera un token nuevo"),
        ("'Probar FB' sale error 200",
         "Usaste el User Token en lugar del Page Token",
         "Repite el Paso E — copia el token de /me/accounts, no el User Token"),
        ("El post se aprueba pero no se publica",
         "El token no tiene permiso pages_manage_posts",
         "Revisa los permisos en el Explorer y regenera el token"),
        ("LinkedIn falla siempre",
         "Permisos aún no aprobados por LinkedIn",
         "Espera 1-3 días hábiles después de solicitar acceso"),
        ("La generación queda procesando",
         "Error en la API de Anthropic o timeout",
         "Actualiza la página. El error aparece en rojo si falló"),
        ("La marca no aparece en el panel",
         "El servidor no estaba corriendo al guardar",
         "Actualiza la página con F5 y vuelve a intentar"),
        ("El logo no se muestra",
         "La marca no estaba guardada antes de subir el logo",
         "Guarda la marca primero, luego sube el logo"),
        ("El análisis con IA no genera nada",
         "El sitio web no es accesible o bloqueó el acceso",
         "Intenta con el manual de marca en su lugar"),
    ],
    col_widths=[4.5, 5.5, 7]
)

# ══════════════════════════════════════════════════════════════════════════
#  PIE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Voxify · Guía generada automáticamente · Junio 2026")
r.font.size = Pt(9); r.font.color.rgb = GRAY; r.italic = True

# ── Guardar ───────────────────────────────────────────────────────────────
output = r"C:\Users\yaco8\OneDrive\Documentos\Voxify - Claude\Voxify_Guia_Configuracion_Marca.docx"
doc.save(output)
print("Documento guardado en:")
print(output)
